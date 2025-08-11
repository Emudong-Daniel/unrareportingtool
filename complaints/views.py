# complaints/views.py
import csv
import io
from statistics import median

from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.urls import reverse
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib import messages
from django.conf import settings
from django.core.mail import send_mail
from django.db.models import Q, Count, F, ExpressionWrapper, DurationField
from django.db.models.functions import TruncMonth
from django.contrib.auth.models import User

# Twilio import made optional
try:
    from twilio.rest import Client
except Exception:
    Client = None

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document
from docx.enum.section import WD_ORIENT

from .models import Complaint, StatusUpdate
from .forms import ComplaintForm, StatusUpdateForm, LookupForm, ReportForm


def send_sms(to, body):
    sid = getattr(settings, 'TWILIO_SID', '') or ''
    token = getattr(settings, 'TWILIO_AUTH_TOKEN', '') or ''
    from_phone = getattr(settings, 'TWILIO_PHONE', '') or ''
    if not (Client and sid and token and from_phone and to):
        return
    try:
        client = Client(sid, token)
        client.messages.create(body=body, from_=from_phone, to=to)
    except Exception:
        return


def is_manager(user):
    return user.is_superuser or user.groups.filter(name='Manager').exists()


def is_technician(user):
    return user.groups.filter(name='Technician').exists()


class RoleLoginView(LoginView):
    template_name = 'complaints/login.html'

    def get_success_url(self):
        user = self.request.user
        if is_technician(user):
            return reverse('complaints:technician_dashboard')
        if is_manager(user):
            return reverse('complaints:dashboard')
        return reverse('complaints:create')


class RoleLogoutView(LogoutView):
    next_page = 'complaints:create'


@login_required
def role_redirect(request):
    user = request.user
    if is_manager(user):
        return redirect('complaints:dashboard')
    if is_technician(user):
        return redirect('complaints:technician_dashboard')
    return redirect('complaints:status_lookup')


def complaint_create(request):
    if request.user.is_authenticated and (is_manager(request.user) or is_technician(request.user)):
        return role_redirect(request)

    if request.method == 'POST':
        form = ComplaintForm(request.POST, request.FILES)
        if form.is_valid():
            comp = form.save()
            if comp.email:
                lookup_url = request.build_absolute_uri(reverse('complaints:status_lookup'))
                send_mail(
                    'Complaint Received',
                    f'Your complaint has been received.\n\n'
                    f'Complaint ID: {comp.id}\n'
                    f'Check status: {lookup_url}',
                    settings.DEFAULT_FROM_EMAIL,
                    [comp.email],
                    fail_silently=False,
                )
            if comp.contact:
                send_sms(
                    comp.contact,
                    f'Complaint #{comp.id} received. Use your ID to check status online.'
                )
            messages.success(request, 'Complaint submitted successfully.')
            return redirect('complaints:complaint_submitted', pk=comp.id)
    else:
        form = ComplaintForm()
    return render(request, 'complaints/complaint_form.html', {'form': form})


def complaint_submitted(request, pk):
    comp = get_object_or_404(Complaint, pk=pk)
    return render(request, 'complaints/complaint_submitted.html', {'complaint': comp})


def status_lookup(request):
    form = LookupForm()
    comp = status = error = None
    if request.method == 'POST':
        form = LookupForm(request.POST)
        if form.is_valid():
            cid = form.cleaned_data['complaint_id']
            contact = form.cleaned_data['contact']
            try:
                comp = Complaint.objects.get(pk=cid, contact=contact)
                status = comp.get_status_display()
            except Complaint.DoesNotExist:
                error = 'No matching complaint found.'
    return render(request, 'complaints/status_lookup.html', {
        'form': form,
        'complaint': comp,
        'status': status,
        'error': error,
    })


@login_required
@user_passes_test(is_manager)
def admin_dashboard(request):
    complaints = Complaint.objects.order_by('status', '-created_at')

    flt = request.GET.get('filter', '').strip()
    if flt == 'closed_by_tech':
        complaints = complaints.filter(status__in=['FIX', 'CLO'], assigned_to__isnull=False)

    query = request.GET.get('q', '').strip()
    if query:
        filters = Q()
        qlow = query.lower()
        matched = [code for code, label in Complaint.STATUS_CHOICES
                   if qlow in code.lower() or qlow in label.lower()]
        if query.isdigit():
            filters |= Q(id=int(query))
        filters |= Q(location__icontains=query)
        filters |= Q(description__icontains=query)
        if matched:
            filters |= Q(status__in=matched)
        complaints = complaints.filter(filters)

    technicians = User.objects.filter(groups__name='Technician')

    if request.method == 'POST':
        cid = request.POST.get('complaint')
        tid = request.POST.get('technician')
        comp = get_object_or_404(Complaint, pk=cid)
        if comp.status in ['FIX', 'CLO']:
            messages.error(request, f"Complaint #{comp.id} is already {comp.get_status_display()} and cannot reassign.")
            return redirect('complaints:admin_dashboard')
        if not (cid and tid):
            messages.error(request, "Select both complaint and technician.")
            return redirect('complaints:admin_dashboard')
        tech = get_object_or_404(User, pk=tid)
        StatusUpdate.objects.create(
            complaint=comp,
            status='INP',
            comment=f'Assigned to {tech.username}'
        )
        comp.assigned_to = tech
        comp.status = 'INP'
        comp.save()
        if comp.email:
            send_mail(
                'Complaint Assignment',
                f'Your complaint #{comp.id} has been assigned to {tech.username}.',
                settings.DEFAULT_FROM_EMAIL,
                [comp.email],
                fail_silently=False,
            )
        messages.success(request, f'Complaint #{comp.id} assigned to {tech.username}.')
        return redirect('complaints:admin_dashboard')

    return render(request, 'complaints/admin_dashboard.html', {
        'complaints': complaints,
        'technicians': technicians,
        'query': query,
    })


@login_required
@user_passes_test(lambda u: is_manager(u) or is_technician(u))
def complaint_update(request, pk):
    comp = get_object_or_404(Complaint, pk=pk)
    if comp.status in ['FIX', 'CLO']:
        messages.error(request, f"Complaint #{comp.id} is already {comp.get_status_display()} and cannot be updated.")
        return redirect('complaints:admin_dashboard')
    if request.method == 'POST':
        form = StatusUpdateForm(request.POST)
        if form.is_valid():
            upd = form.save(commit=False)
            upd.complaint = comp
            upd.save()
            comp.status = upd.status
            comp.save()
            if comp.email:
                send_mail(
                    'Complaint Status Update',
                    f'Your complaint #{comp.id} status is now {comp.get_status_display()}.',
                    settings.DEFAULT_FROM_EMAIL,
                    [comp.email],
                    fail_silently=False,
                )
            messages.success(request, 'Status updated successfully.')
            if is_manager(request.user):
                return redirect('complaints:admin_dashboard')
            return redirect('complaints:complaint_update', pk=pk)
    else:
        form = StatusUpdateForm()
    return render(request, 'complaints/complaint_update.html', {
        'form': form,
        'complaint': comp,
    })


@login_required
@user_passes_test(is_technician)
def technician_dashboard(request):
    qs = Complaint.objects.filter(assigned_to=request.user).order_by('-created_at')

    total = qs.count()
    status_counts = {
        'new': qs.filter(status='NEW').count(),
        'in_progress': qs.filter(status='INP').count(),
        'fixed': qs.filter(status='FIX').count(),
        'closed': qs.filter(status='CLO').count(),
    }

    resolved_deltas = qs.filter(status='CLO') \
        .annotate(resolution=ExpressionWrapper(
            F('updates__timestamp') - F('created_at'),
            output_field=DurationField()
        )) \
        .values_list('resolution', flat=True)

    days = [delta.total_seconds() / 86400 for delta in resolved_deltas if delta]
    avg_resolution = round(sum(days) / len(days), 2) if days else 0
    median_resolution = round(median(days), 2) if days else 0

    return render(request, 'complaints/technician_dashboard.html', {
        'complaints': qs,
        'total': total,
        'status_counts': status_counts,
        'avg_resolution': avg_resolution,
        'median_resolution': median_resolution,
    })


@login_required
@user_passes_test(is_manager)
def dashboard(request):
    """
    Metrics. Closed includes FIX and CLO that were assigned to a technician.
    Monthly rows show month names.
    """
    total = Complaint.objects.count()

    status_qs = Complaint.objects.values('status').annotate(cnt=Count('id'))
    status_counts = {s['status']: s['cnt'] for s in status_qs}

    closed_by_tech = Complaint.objects.filter(status__in=['FIX', 'CLO'], assigned_to__isnull=False).count()

    breakdown = {
        'new': status_counts.get('NEW', 0),
        'in_progress': status_counts.get('INP', 0),
        'fixed': status_counts.get('FIX', 0),
        'closed': closed_by_tech,
    }

    resolved = Complaint.objects.filter(status__in=['FIX', 'CLO']).annotate(
        resolution=ExpressionWrapper(
            F('updates__timestamp') - F('created_at'),
            output_field=DurationField()
        )
    ).values_list('resolution', flat=True)
    days = [td.total_seconds() / 86400 for td in resolved if td]
    avg_resolution = round(sum(days) / len(days), 2) if days else 0
    median_resolution = round(median(days), 2) if days else 0

    # Month names using TruncMonth
    monthly = (
        Complaint.objects
        .annotate(month=TruncMonth('created_at'))
        .values('month')
        .annotate(count=Count('id'))
        .order_by('month')
    )

    top_locations = (
        Complaint.objects
        .values('location')
        .annotate(count=Count('id'))
        .order_by('-count')[:5]
    )

    return render(request, 'complaints/dashboard.html', {
        'total': total,
        'breakdown': breakdown,
        'avg_resolution': avg_resolution,
        'median_resolution': median_resolution,
        'monthly': list(monthly),
        'top_locations': top_locations,
    })


@login_required
@user_passes_test(is_manager)
def reports(request):
    form = ReportForm(request.GET or None)
    qs = Complaint.objects.order_by('-created_at')
    if form.is_valid():
        sd = form.cleaned_data['start_date']
        ed = form.cleaned_data['end_date']
        st = form.cleaned_data['status']
        loc = form.cleaned_data['location']
        if sd:
            qs = qs.filter(created_at__date__gte=sd)
        if ed:
            qs = qs.filter(created_at__date__lte=ed)
        if st:
            qs = qs.filter(status=st)
        if loc:
            qs = qs.filter(location__icontains=loc)

    export = request.GET.get('export')
    if export == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="complaints_report.csv"'
        writer = csv.writer(response)
        writer.writerow(['ID', 'Name', 'Contact', 'Email', 'Location', 'Status', 'Created', 'Description'])
        for c in qs:
            writer.writerow([
                c.id, c.name or '', c.contact or '', c.email or '',
                c.location, c.get_status_display(), c.created_at.strftime('%Y-%m-%d'), c.description
            ])
        return response

    if export == 'pdf':
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
        styles = getSampleStyleSheet()
        elements = [Paragraph('Complaints Report', styles['Title'])]
        data = [['ID', 'Name', 'Contact', 'Email', 'Location', 'Status', 'Created', 'Description']]
        for c in qs:
            data.append([
                str(c.id), c.name or '', c.contact or '', c.email or '',
                c.location, c.get_status_display(), c.created_at.strftime('%Y-%m-%d'), c.description
            ])
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ]))
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="complaints_report.pdf"'
        return response

    if export == 'word':
        docx = Document()
        section = docx.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        docx.add_heading('Complaints Report', 0)
        table = docx.add_table(rows=1, cols=8)
        hdr = table.rows[0].cells
        for i, h in enumerate(['ID', 'Name', 'Contact', 'Email', 'Location', 'Status', 'Created', 'Description']):
            hdr[i].text = h
        for c in qs:
            row = table.add_row().cells
            row[0].text = str(c.id)
            row[1].text = c.name or ''
            row[2].text = c.contact or ''
            row[3].text = c.email or ''
            row[4].text = c.location
            row[5].text = c.get_status_display()
            row[6].text = c.created_at.strftime('%Y-%m-%d')
            row[7].text = c.description
        buf = io.BytesIO()
        docx.save(buf)
        buf.seek(0)
        response = HttpResponse(
            buf.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="complaints_report.docx"'
        return response

    return render(request, 'complaints/reports.html', {
        'form': form,
        'complaints': qs,
    })
